"""
Resume Parser Module

This module extracts information from resumes in PDF and DOCX formats,
including skills, education, experience, and contact information.
"""

import os
import re
import logging
from typing import Dict, List, Set, Any, Optional
from datetime import datetime

# PDF parsing
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not available")

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

# Advanced PDF parsing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parse resumes from PDF and DOCX files to extract structured information.
    """
    
    def __init__(self):
        """Initialize the resume parser."""
        self.skills_database = self._load_skills_database()
        self.education_keywords = self._load_education_keywords()
        self.experience_keywords = self._load_experience_keywords()
    
    def _load_skills_database(self) -> Set[str]:
        """Load comprehensive skills database."""
        skills = {
            # Programming Languages
            'python', 'java', 'javascript', 'c++', 'c#', 'c', 'php', 'ruby', 'go', 'rust',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'typescript', 'dart', 'lua',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue.js', 'vue', 'node.js', 'express', 'django',
            'flask', 'spring', 'spring boot', 'asp.net', 'jquery', 'bootstrap', 'tailwind',
            'next.js', 'nuxt.js', 'svelte', 'gatsby', 'webpack', 'babel', 'sass', 'less',
            
            # Databases
            'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server', 'sqlite',
            'redis', 'cassandra', 'dynamodb', 'elasticsearch', 'neo4j', 'mariadb',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins',
            'gitlab ci', 'github actions', 'terraform', 'ansible', 'chef', 'puppet',
            'ci/cd', 'devops', 'cloud', 'serverless', 'lambda', 'ec2', 's3',
            
            # Data Science & ML
            'machine learning', 'ml', 'deep learning', 'ai', 'artificial intelligence',
            'neural networks', 'nlp', 'computer vision', 'tensorflow', 'pytorch', 'keras',
            'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly',
            'data science', 'data analysis', 'statistics', 'big data', 'spark', 'hadoop',
            
            # Tools & Frameworks
            'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
            'visual studio', 'vs code', 'intellij', 'eclipse', 'pycharm', 'jupyter',
            'postman', 'swagger', 'graphql', 'rest api', 'soap', 'microservices',
            
            # Mobile Development
            'android', 'ios', 'react native', 'flutter', 'xamarin', 'kotlin', 'swift',
            'objective-c', 'mobile development', 'app development',
            
            # Other Technologies
            'agile', 'scrum', 'kanban', 'test automation', 'selenium', 'cypress',
            'unit testing', 'integration testing', 'tdd', 'bdd', 'linux', 'unix',
            'bash', 'powershell', 'security', 'networking', 'api', 'restful',
            
            # Business & Soft Skills
            'leadership', 'management', 'communication', 'problem solving', 'teamwork',
            'project management', 'agile methodologies', 'business analysis',
            'stakeholder management', 'requirements gathering', 'documentation',
            
            # Data Visualization & BI
            'tableau', 'power bi', 'looker', 'qlik', 'excel', 'google sheets',
            'data visualization', 'business intelligence', 'reporting',
            
            # Blockchain & Emerging Tech
            'blockchain', 'solidity', 'ethereum', 'smart contracts', 'web3', 'nft',
            'cryptocurrency', 'iot', 'ar', 'vr', 'augmented reality', 'virtual reality'
        }
        
        return skills
    
    def _load_education_keywords(self) -> List[str]:
        """Load education keywords."""
        return [
            'education', 'qualification', 'academic', 'degree', 'bachelor', 'master',
            'phd', 'doctorate', 'diploma', 'certificate', 'b.tech', 'm.tech', 'b.e',
            'm.e', 'b.sc', 'm.sc', 'mba', 'bba', 'bca', 'mca', 'university', 'college',
            'institute', 'school', 'graduated', 'graduation', 'undergraduate', 'postgraduate'
        ]
    
    def _load_experience_keywords(self) -> List[str]:
        """Load experience keywords."""
        return [
            'experience', 'work experience', 'employment', 'work history', 'career',
            'professional experience', 'job', 'position', 'role', 'responsibilities',
            'worked', 'working', 'employed', 'internship', 'intern', 'trainee'
        ]
    
    def parse_pdf_pypdf2(self, file_path: str) -> str:
        """
        Extract text from PDF using PyPDF2.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 is not installed")
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error parsing PDF with PyPDF2: {e}")
            raise
        
        return text
    
    def parse_pdf_pdfplumber(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber (more accurate).
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
        """
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is not installed")
        
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error parsing PDF with pdfplumber: {e}")
            raise
        
        return text
    
    def parse_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed")
        
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise
        
        return text
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from resume file (auto-detect format).
        
        Args:
            file_path (str): Path to resume file
            
        Returns:
            str: Extracted text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # Try pdfplumber first (more accurate), fallback to PyPDF2
            try:
                if PDFPLUMBER_AVAILABLE:
                    return self.parse_pdf_pdfplumber(file_path)
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            if PYPDF2_AVAILABLE:
                return self.parse_pdf_pypdf2(file_path)
            else:
                raise ImportError("No PDF parsing library available")
        
        elif file_ext in ['.docx', '.doc']:
            return self.parse_docx(file_path)
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        # Pattern for Indian phone numbers
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
        matches = re.findall(phone_pattern, text)
        
        # Filter out years and other numbers
        for match in matches:
            cleaned = re.sub(r'[-\s\.\(\)]', '', match)
            if 10 <= len(cleaned) <= 15:
                return match
        
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from resume text.
        
        Args:
            text (str): Resume text
            
        Returns:
            List[str]: List of identified skills
        """
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.skills_database:
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill.title())
        
        # Remove duplicates and sort
        found_skills = sorted(list(set(found_skills)))
        
        return found_skills
    
    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """
        Extract education information.
        
        Args:
            text (str): Resume text
            
        Returns:
            List[Dict]: List of education entries
        """
        education = []
        lines = text.split('\n')
        
        degrees = {
            'bachelor': ['bachelor', 'b.tech', 'b.e', 'b.sc', 'bba', 'bca', 'ba', 'bcom'],
            'master': ['master', 'm.tech', 'm.e', 'm.sc', 'mba', 'mca', 'ma', 'mcom'],
            'phd': ['phd', 'ph.d', 'doctorate', 'doctor of philosophy'],
            'diploma': ['diploma', 'certificate']
        }
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            for degree_type, keywords in degrees.items():
                if any(keyword in line_lower for keyword in keywords):
                    # Try to find year
                    year_match = re.search(r'20[0-2][0-9]|19[89][0-9]', line)
                    year = year_match.group() if year_match else 'N/A'
                    
                    # Try to find institution (next few lines)
                    institution = 'N/A'
                    for j in range(i, min(i+3, len(lines))):
                        if any(keyword in lines[j].lower() for keyword in ['university', 'college', 'institute']):
                            institution = lines[j].strip()
                            break
                    
                    education.append({
                        'degree': degree_type.title(),
                        'detail': line.strip(),
                        'institution': institution,
                        'year': year
                    })
                    break
        
        return education
    
    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        """
        Extract work experience information.
        
        Args:
            text (str): Resume text
            
        Returns:
            List[Dict]: List of experience entries
        """
        experience = []
        lines = text.split('\n')
        
        # Common job title patterns
        job_titles = [
            'engineer', 'developer', 'analyst', 'manager', 'consultant', 'designer',
            'architect', 'lead', 'senior', 'junior', 'intern', 'specialist', 'expert',
            'director', 'coordinator', 'administrator', 'assistant', 'associate'
        ]
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Check if line contains job title
            if any(title in line_lower for title in job_titles):
                # Try to find duration
                duration_match = re.search(r'(20[0-2][0-9]|19[89][0-9])\s*-\s*(20[0-2][0-9]|19[89][0-9]|present|current)', line_lower)
                duration = duration_match.group() if duration_match else 'N/A'
                
                # Try to find company (next line might be company)
                company = lines[i+1].strip() if i+1 < len(lines) else 'N/A'
                
                experience.append({
                    'title': line.strip(),
                    'company': company,
                    'duration': duration
                })
        
        return experience[:10]  # Limit to 10 entries
    
    def extract_years_of_experience(self, text: str) -> Optional[int]:
        """
        Estimate total years of experience.
        
        Args:
            text (str): Resume text
            
        Returns:
            Optional[int]: Estimated years of experience
        """
        # Look for explicit mentions
        patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'experience\s*[:=]\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*experience'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text.lower())
            if match:
                return int(match.group(1))
        
        # Estimate from experience dates
        years = re.findall(r'20[0-2][0-9]|19[89][0-9]', text)
        if len(years) >= 2:
            years = [int(y) for y in years]
            return max(years) - min(years)
        
        return None
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume and extract all information.
        
        Args:
            file_path (str): Path to resume file
            
        Returns:
            Dict[str, Any]: Extracted resume information
        """
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            # Extract all information
            result = {
                'file_name': os.path.basename(file_path),
                'parsed_at': datetime.now().isoformat(),
                'contact': {
                    'email': self.extract_email(text),
                    'phone': self.extract_phone(text)
                },
                'skills': self.extract_skills(text),
                'education': self.extract_education(text),
                'experience': self.extract_experience(text),
                'years_of_experience': self.extract_years_of_experience(text),
                'raw_text_preview': text[:500] + '...' if len(text) > 500 else text
            }
            
            logger.info(f"Successfully parsed resume: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {e}")
            raise
    
    def calculate_resume_score(self, resume_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Calculate resume quality score.
        
        Args:
            resume_data (Dict): Parsed resume data
            
        Returns:
            Dict: Score breakdown
        """
        score = 0
        max_score = 100
        feedback = []
        
        # Contact information (20 points)
        if resume_data['contact']['email']:
            score += 10
        else:
            feedback.append("Add email address")
        
        if resume_data['contact']['phone']:
            score += 10
        else:
            feedback.append("Add phone number")
        
        # Skills (30 points)
        skills_count = len(resume_data['skills'])
        if skills_count >= 10:
            score += 30
        elif skills_count >= 5:
            score += 20
            feedback.append("Add more relevant skills")
        else:
            score += 10
            feedback.append("Add technical skills")
        
        # Education (20 points)
        if resume_data['education']:
            score += 20
        else:
            feedback.append("Add education details")
        
        # Experience (30 points)
        if resume_data['experience']:
            exp_count = len(resume_data['experience'])
            if exp_count >= 3:
                score += 30
            elif exp_count >= 1:
                score += 20
            else:
                score += 10
        else:
            feedback.append("Add work experience")
        
        # Calculate percentage
        percentage = (score / max_score) * 100
        
        # Determine grade
        if percentage >= 80:
            grade = 'A - Excellent'
        elif percentage >= 60:
            grade = 'B - Good'
        elif percentage >= 40:
            grade = 'C - Average'
        else:
            grade = 'D - Needs Improvement'
        
        return {
            'score': score,
            'max_score': max_score,
            'percentage': round(percentage, 2),
            'grade': grade,
            'feedback': feedback
        }


def main():
    """Test resume parser."""
    parser = ResumeParser()
    
    print("="*70)
    print("Resume Parser - Test")
    print("="*70)
    
    # Test text extraction
    sample_text = """
    John Doe
    john.doe@email.com | +91-9876543210
    
    EXPERIENCE
    Senior Software Engineer at TechCorp (2020 - Present)
    Developed web applications using React and Node.js
    
    Software Developer at StartupXYZ (2018 - 2020)
    Worked on Python and Django projects
    
    EDUCATION
    Bachelor of Technology in Computer Science
    XYZ University, 2018
    
    SKILLS
    Python, JavaScript, React, Node.js, SQL, MongoDB, Docker, AWS, Git
    Machine Learning, Data Analysis, REST API
    """
    
    print("\n1. Extracting Email:")
    email = parser.extract_email(sample_text)
    print(f"   Email: {email}")
    
    print("\n2. Extracting Phone:")
    phone = parser.extract_phone(sample_text)
    print(f"   Phone: {phone}")
    
    print("\n3. Extracting Skills:")
    skills = parser.extract_skills(sample_text)
    print(f"   Found {len(skills)} skills: {', '.join(skills[:10])}")
    if len(skills) > 10:
        print(f"   ... and {len(skills) - 10} more")
    
    print("\n4. Extracting Education:")
    education = parser.extract_education(sample_text)
    for edu in education:
        print(f"   - {edu['degree']}: {edu['detail']}")
    
    print("\n5. Extracting Experience:")
    experience = parser.extract_experience(sample_text)
    for exp in experience:
        print(f"   - {exp['title']}")
        print(f"     Company: {exp['company']}")
        print(f"     Duration: {exp['duration']}")
    
    print("\n6. Years of Experience:")
    years = parser.extract_years_of_experience(sample_text)
    print(f"   Estimated: {years} years")
    
    print("\n7. Resume Score:")
    resume_data = {
        'contact': {'email': email, 'phone': phone},
        'skills': skills,
        'education': education,
        'experience': experience
    }
    score_data = parser.calculate_resume_score(resume_data)
    print(f"   Score: {score_data['score']}/{score_data['max_score']} ({score_data['percentage']}%)")
    print(f"   Grade: {score_data['grade']}")
    if score_data['feedback']:
        print(f"   Feedback:")
        for item in score_data['feedback']:
            print(f"      - {item}")
    
    print("\n" + "="*70)
    print("Resume Parser Test Complete!")
    print("="*70)


if __name__ == "__main__":
    main()

